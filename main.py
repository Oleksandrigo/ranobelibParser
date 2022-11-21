import os
import time

from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from htmldocx import HtmlToDocx

from selenium.common import exceptions as sel_exeptions
from selenium.webdriver import ChromeOptions, DesiredCapabilities
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.webdriver.common.by import By

from webdriver_manager.chrome import ChromeDriverManager
from undetected_chromedriver import Chrome as Chrome_UC

import cloudscraper
from requests import Response
from requests_html import HTMLSession, HTMLResponse, HTML

BASE_URL = "https://ranobelib.me"


def create_browser() -> WebDriver:
    options = ChromeOptions()
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--no-sandbox")
    service = Service(ChromeDriverManager().install())
    capa = DesiredCapabilities.CHROME
    capa["pageLoadStrategy"] = "none"
    browser = Chrome_UC(options=options, service=service, desired_capabilities=capa)
    browser.delete_all_cookies()
    return browser


def get_title_url(link: str) -> str:
    title_url = ""
    if BASE_URL in link:
        s = link.split(BASE_URL)[1]
        if len(s.split("/")) > 2:
            title_url = "/" + s.split("/")[1]
        if "?" in s:
            title_url = s.split("?")[0].strip()

    return title_url


def get_all_chapters(browser: WebDriver) -> list[tuple[str, str]]:
    print("В зависимости от количества глав, сейчас может слегка зависнуть. Ничего не трогайте!")
    chapters_vars = ["/html/body/div[1]/div[3]/div/div[3]/div[2]",
                     "/html/body/div[1]/div[3]/div/div[2]/div/div[1]"]
    try:
        chapters_div = browser.find_element(by=By.XPATH, value=chapters_vars[0])
        chapters_div.click()

    except Exception:
        chapters_div = browser.find_element(by=By.XPATH, value=chapters_vars[1])
        chapters_div.click()

    time.sleep(2)
    data = HTML(html=browser.page_source)
    print("Браузер закрывается, так и должно быть.")

    browser.quit()

    chapters_divs = data.find("div.popup__content")
    try:
        chapters_divs = chapters_divs[0].find("div.modal__body")[0].find("a")
    except Exception:
        print("Не найден список глав на странице. БААААГ")
        quit()

    chapters = []
    for chapt in chapters_divs[::-1]:
        if not chapt.text.strip() or not chapt.attrs.get("href"):
            print("У названия главы на найден текст или ссылка.\n"
                  "Попробуйте снова.")
            quit()

        chapters.append((chapt.text, BASE_URL + chapt.attrs.get("href")))
    return chapters


def parse_chapers(name_title: str, chapters: list[tuple[str, str]]):
    # print("Создание скрапера CloudFire")
    # scraper = cloudscraper.create_scraper()
    # print("Получение HTML")
    # r = scraper.get(url)
    # print("Запихивается в requsts-html")
    # res = HTML(html=r.text)
    browser_args = [
        "--no-sandbox",
        "--disable-blink-features=AutomationControlled",
        "--ignore-certificate-errors"
    ]
    session = HTMLSession(browser_args=browser_args)
    doc_file = Document()
    new_parser = HtmlToDocx()

    for num, (name, url_chapt) in enumerate(chapters):
        os.system('cls||clear')
        print(f"{num + 1} / {len(chapters)}")
        chapter: HTMLResponse | Response = session.get(url_chapt)
        if chapter.status_code == 200:
            datas = chapter.html.find("div.reader-container.container.container_center")

            title_head = doc_file.add_heading(name, 1)
            title_head.style.font.color.rgb = RGBColor.from_string("000000")
            new_parser.add_html_to_document(datas[0].html, doc_file)
        else:
            print("Произошла ошибка, страница скорее всего заблочена из-за слишком частых запросов.")
            tries = 5
            while tries >= 0:
                print("")
                print("Пробуем дождаться ответа страницы. ждем 10 секунд")
                time.sleep(10)
                chapter: HTMLResponse | Response = session.get(url_chapt)
                if chapter.status_code != 200:
                    tries -= 1
                    continue
                else:
                    datas = chapter.html.find("div.reader-container.container.container_center")

                    title_head: Paragraph = doc_file.add_heading(name, 1)
                    title_head.style.font.color.rgb = RGBColor.from_string("000000")
                    new_parser.add_html_to_document(datas[0].html, doc_file)
                    break

    doc_file.save(f'{name_title}.docx')


def start_parse(link: str):
    time_wait = 5.0
    title_url = get_title_url(link)
    url = BASE_URL + title_url
    browser: WebDriver = create_browser()

    print("Запущен браузер. (chromedriver)")
    print("Нужен только до получения списка глав, дальше закроется сам.")
    print("Не трогайте, не двигайте, не меняйте размеры, не сворачивайте. "
          "НИКАК НЕ ВЗАИМОДЕЙСТВУЙТЕ С ЭТИМ ОКНОМ!!")
    print("ПОКА ЭТО ОКНО САМО НЕ ЗАКРОЕТСЯ НЕ ПЕРЕКЛЮЧАЙТЕСЬ НА ДРУГИЕ ОКНА!!!!!!!")
    print("Открывается ссылка, введенная пользователем")
    browser.get(url)
    time.sleep(time_wait)
    main_title = browser.find_element(by=By.XPATH,
                                      value="/html/body/div[3]/div/div/div/div[2]/div[1]/div[1]/div[1]").text
    alt_title = browser.find_element(by=By.XPATH,
                                     value="/html/body/div[3]/div/div/div/div[2]/div[1]/div[1]/div[2]").text
    print(main_title, alt_title)
    print("Переходим на страницу первой главы")
    browser.find_element(by=By.LINK_TEXT, value="Начать читать").click()
    time.sleep(time_wait)
    try:
        browser.find_element(by=By.XPATH, value="/html/body/div[3]/div/div/div[2]/div/button[2]").click()
        time.sleep(3)
    except Exception:
        pass
    print("Получаем список глав")
    chapters = get_all_chapters(browser)
    print("Начинаем парсинг глав!")
    parse_chapers(main_title, chapters)
    print("Завершено!")


if __name__ == "__main__":
    input_url = input("Введите ссылку на ранобе в ranobelib.me для парсинга.\n")
    # input_url = "https://ranobelib.me/longwang-chuan?section=info"
    start_parse(input_url)
